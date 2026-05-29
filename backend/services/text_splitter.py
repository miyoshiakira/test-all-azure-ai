"""
RAG向けテキストスプリッター - 意味分割によるチャンク化

対応ファイル形式:
- PDF: 見出し構造を検出してセクション単位で分割
- DOCX: 見出しスタイルを認識してセクション単位で分割
- Excel: シート＋ヘッダー行を維持して行グループ単位で分割
- CSV: ヘッダー行を維持して行グループ単位で分割

意味分割の考え方:
- 単純な文字数分割ではなく、文書の構造（見出し・段落・表）を維持
- 各チャンクが1つの意味的トピックを含むように分割
- チャンク間のオーバーラップで文脈を維持
"""
import io
import csv
import re
from typing import List, Tuple, Optional
from dataclasses import dataclass, field


@dataclass
class Chunk:
    """チャンク（意味的に分割されたテキスト単位）"""
    text: str
    chunk_id: str
    chunk_type: str  # "section", "page_section", "sheet_rows", "csv_rows"
    metadata: dict = field(default_factory=dict)  # ページ番号・シート名などのメタ情報


class TextSplitter:
    """RAG向け意味分割テキストスプリッター"""

    # チャンク設定
    MAX_CHUNK_SIZE = 800       # チャンクの最大文字数（埋め込みモデルのトークン制限を考慮）
    MIN_CHUNK_SIZE = 100       # チャンクの最小文字数（これ未満は前チャンクに結合）
    OVERLAP_SIZE = 100         # チャンク間のオーバーラップ文字数
    HEADING_PATTERNS = [
        r'^#{1,6}\s',           # Markdown: # 見出し
        r'^■\s?',               # 日本語見出し: ■
        r'^●\s?',               # 日本語見出し: ●
        r'^◆\s?',               # 日本語見出し: ◆
        r'^【.+?】',             # 日本語見出し: 【】
        r'^\d+[\.．、)\）]\s',   # 番号付き見出し: 1. / 1、/ 1)
        r'^[０-９]+[\.．、)\）]\s',  # 全角番号付き見出し
    ]

    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.xlsx', '.xls', '.csv'}

    @staticmethod
    def is_supported(file_name: str) -> bool:
        """対応ファイル形式か判定"""
        ext = '.' + file_name.rsplit('.', 1)[-1].lower() if '.' in file_name else ''
        return ext in TextSplitter.SUPPORTED_EXTENSIONS

    @staticmethod
    def split(file_content: bytes, file_name: str) -> Tuple[List[Chunk], str]:
        """
        ファイルを読み込み、意味的にチャンク分割
        Returns: (chunks, file_type)
        """
        ext = '.' + file_name.rsplit('.', 1)[-1].lower() if '.' in file_name else ''

        if ext == '.pdf':
            return TextSplitter._split_pdf(file_content), "PDF"
        elif ext in ('.docx',):
            return TextSplitter._split_docx(file_content), "Word"
        elif ext in ('.xlsx', '.xls'):
            return TextSplitter._split_xlsx(file_content), "Excel"
        elif ext == '.csv':
            return TextSplitter._split_csv(file_content), "CSV"
        else:
            return [Chunk(text=f"[未対応のファイル形式: {file_name}]", chunk_id="error", chunk_type="error")], "Unknown"

    # ============================================================
    # PDF
    # ============================================================
    @staticmethod
    def _split_pdf(content: bytes) -> List[Chunk]:
        """PDFから見出し構造を検出してセクション単位でチャンク分割"""
        import fitz  # PyMuPDF

        chunks = []
        with fitz.open(stream=content, filetype="pdf") as doc:
            total_pages = len(doc)

            # 全ページのテキストを行単位で収集
            all_lines = []
            for page_num, page in enumerate(doc, 1):
                page_text = page.get_text()
                for line in page_text.split('\n'):
                    line = line.strip()
                    if line:
                        all_lines.append((page_num, line))

            if not all_lines:
                return [Chunk(text="[PDFからテキストを抽出できませんでした]", chunk_id="error", chunk_type="error")]

            # 見出し位置を検出
            heading_indices = []
            for i, (page_num, line) in enumerate(all_lines):
                if TextSplitter._is_heading(line):
                    heading_indices.append(i)

            # 見出しがなければ、ページ境界＋トピック分割
            if not heading_indices:
                # ページごとにテキストをまとめ、意味分割
                page_texts = {}
                for page_num, line in all_lines:
                    if page_num not in page_texts:
                        page_texts[page_num] = []
                    page_texts[page_num].append(line)

                # ページテキストを結合してトピック分割
                full_sections = []
                for page_num in sorted(page_texts.keys()):
                    text = '\n'.join(page_texts[page_num])
                    full_sections.append((page_num, text))

                # トピックベースでチャンク化
                chunks = TextSplitter._build_chunks_from_sections(
                    full_sections, chunk_type="page_section"
                )
            else:
                # 見出し位置でセクションに分割
                sections = []
                for sec_idx, start in enumerate(heading_indices):
                    end = heading_indices[sec_idx + 1] if sec_idx + 1 < len(heading_indices) else len(all_lines)
                    sec_lines = all_lines[start:end]
                    page_num = sec_lines[0][0]
                    sec_text = '\n'.join(line for _, line in sec_lines)
                    sections.append((page_num, sec_text))

                # 最初の見出しより前のテキスト
                if heading_indices[0] > 0:
                    pre_lines = all_lines[:heading_indices[0]]
                    pre_text = '\n'.join(line for _, line in pre_lines)
                    sections.insert(0, (pre_lines[0][0], pre_text))

                chunks = TextSplitter._build_chunks_from_sections(
                    sections, chunk_type="section"
                )

        return chunks if chunks else [Chunk(text="[PDFからテキストを抽出できませんでした]", chunk_id="error", chunk_type="error")]

    # ============================================================
    # DOCX
    # ============================================================
    @staticmethod
    def _split_docx(content: bytes) -> List[Chunk]:
        """Word文書を見出しスタイルでセクション分割"""
        from docx import Document
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        doc = Document(io.BytesIO(content))

        # 段落を走査して見出し位置を検出
        elements = []  # (type, text) - type: "heading" | "paragraph" | "table"

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = (para.style.name or "").lower() if para.style else ""
            if 'heading' in style_name or '見出し' in style_name:
                elements.append(("heading", text))
            else:
                elements.append(("paragraph", text))

        # テーブルも追加
        for table in doc.tables:
            table_lines = []
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_lines.append(row_text)
            if table_lines:
                elements.append(("table", '\n'.join(table_lines)))

        if not elements:
            return [Chunk(text="[Wordからテキストを抽出できませんでした]", chunk_id="error", chunk_type="error")]

        # 見出しでセクションに分割
        sections = []
        current_heading = ""
        current_texts = []

        for elem_type, text in elements:
            if elem_type == "heading":
                # 前のセクションを保存
                if current_texts:
                    sections.append((current_heading, '\n\n'.join(current_texts)))
                current_heading = text
                current_texts = []
            else:
                current_texts.append(text)

        # 最後のセクション
        if current_texts:
            sections.append((current_heading, '\n\n'.join(current_texts)))

        # セクションからチャンク構築
        chunks = []
        for sec_idx, (heading, text) in enumerate(sections, 1):
            header = f"[{heading}]\n\n" if heading else ""
            full_text = header + text

            # セクションが大きすぎる場合はトピック分割
            if len(full_text) > TextSplitter.MAX_CHUNK_SIZE:
                sub_chunks = TextSplitter._split_text_semantic(
                    full_text,
                    chunk_id_prefix=f"section_{sec_idx}",
                    chunk_type="section_part"
                )
                chunks.extend(sub_chunks)
            else:
                chunks.append(Chunk(
                    text=full_text,
                    chunk_id=f"section_{sec_idx}",
                    chunk_type="section",
                    metadata={"heading": heading} if heading else {}
                ))

        # 小さすぎるチャンクを結合
        chunks = TextSplitter._merge_small_chunks(chunks)

        return chunks if chunks else [Chunk(text="[Wordからテキストを抽出できませんでした]", chunk_id="error", chunk_type="error")]

    # ============================================================
    # Excel
    # ============================================================
    @staticmethod
    def _split_xlsx(content: bytes) -> List[Chunk]:
        """Excelをシート単位で、ヘッダー行を維持して行グループに分割"""
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        chunks = []

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            rows = []
            for row in sheet.iter_rows():
                row_values = [str(cell.value) if cell.value is not None else "" for cell in row]
                if any(v for v in row_values):  # 空行スキップ
                    rows.append(row_values)

            if not rows:
                continue

            # 1行目をヘッダーとして扱う
            header_row = rows[0]
            header_text = " | ".join(header_row)

            # ヘッダー付きで行をテキスト化
            data_rows = rows[1:] if len(rows) > 1 else []
            if not data_rows:
                # ヘッダーのみ
                chunks.append(Chunk(
                    text=f"[シート: {sheet_name}]\n{header_text}",
                    chunk_id=f"sheet_{sheet_name}",
                    chunk_type="sheet",
                    metadata={"sheet": sheet_name, "row_count": 1}
                ))
                continue

            # 行グループ単位でチャンク化（ヘッダーを各チャンクに付与）
            chunk_rows = []
            current_size = len(header_text) + len(f"[シート: {sheet_name}]\n")

            for row_idx, row in enumerate(data_rows, 2):
                row_text = " | ".join(row)
                row_size = len(row_text) + 1  # +1 for newline

                if current_size + row_size > TextSplitter.MAX_CHUNK_SIZE and chunk_rows:
                    # 現在のチャンクを保存
                    chunk_text = TextSplitter._format_table_chunk(
                        sheet_name, header_text, chunk_rows,
                        start_row=2, end_row=2 + len(chunk_rows) - 1
                    )
                    chunks.append(Chunk(
                        text=chunk_text,
                        chunk_id=f"sheet_{sheet_name}_rows_{2}-{2 + len(chunk_rows) - 1}",
                        chunk_type="sheet_rows",
                        metadata={"sheet": sheet_name, "start_row": 2, "end_row": 2 + len(chunk_rows) - 1}
                    ))
                    chunk_rows = [row]
                    current_size = len(header_text) + row_size + len(f"[シート: {sheet_name}]\n")
                else:
                    chunk_rows.append(row)
                    current_size += row_size

            # 残りの行
            if chunk_rows:
                start = 2 + len(data_rows) - len(chunk_rows)
                chunk_text = TextSplitter._format_table_chunk(
                    sheet_name, header_text, chunk_rows,
                    start_row=start, end_row=start + len(chunk_rows) - 1
                )
                chunks.append(Chunk(
                    text=chunk_text,
                    chunk_id=f"sheet_{sheet_name}_rows_{start}-{start + len(chunk_rows) - 1}",
                    chunk_type="sheet_rows",
                    metadata={"sheet": sheet_name, "start_row": start, "end_row": start + len(chunk_rows) - 1}
                ))

        wb.close()
        return chunks if chunks else [Chunk(text="[Excelからテキストを抽出できませんでした]", chunk_id="error", chunk_type="error")]

    @staticmethod
    def _format_table_chunk(sheet_name: str, header_text: str, rows: List[List[str]], start_row: int, end_row: int) -> str:
        """テーブルチャンクをフォーマット（ヘッダー行を各チャンクに含める）"""
        lines = [f"[シート: {sheet_name} / 行 {start_row}-{end_row}]"]
        lines.append(header_text)
        lines.append("-" * min(len(header_text), 60))
        for row in rows:
            lines.append(" | ".join(row))
        return '\n'.join(lines)

    # ============================================================
    # CSV
    # ============================================================
    @staticmethod
    def _split_csv(content: bytes) -> List[Chunk]:
        """CSVをヘッダー行維持で行グループに分割"""
        # エンコーディング検出（UTF-8 → Shift-JIS → CP932）
        text = None
        for encoding in ['utf-8-sig', 'utf-8', 'shift-jis', 'cp932']:
            try:
                text = content.decode(encoding)
                break
            except (UnicodeDecodeError, LookupError):
                continue

        if text is None:
            text = content.decode('utf-8', errors='ignore')

        reader = csv.reader(io.StringIO(text))
        rows = list(reader)

        if not rows:
            return [Chunk(text="[CSVからテキストを抽出できませんでした]", chunk_id="error", chunk_type="error")]

        header_row = rows[0]
        header_text = ",".join(header_row)
        data_rows = rows[1:] if len(rows) > 1 else []

        if not data_rows:
            return [Chunk(
                text=f"[CSV / ヘッダー行]\n{header_text}",
                chunk_id="csv_header",
                chunk_type="csv_rows",
                metadata={"row_count": 1}
            )]

        chunks = []
        chunk_rows = []
        current_size = len(header_text) + len("[CSV]\n")

        for row_idx, row in enumerate(data_rows, 2):
            row_text = ",".join(row)
            row_size = len(row_text) + 1

            if current_size + row_size > TextSplitter.MAX_CHUNK_SIZE and chunk_rows:
                chunk_text = TextSplitter._format_csv_chunk(
                    header_text, chunk_rows,
                    start_row=2, end_row=2 + len(chunk_rows) - 1
                )
                chunks.append(Chunk(
                    text=chunk_text,
                    chunk_id=f"csv_rows_{2}-{2 + len(chunk_rows) - 1}",
                    chunk_type="csv_rows",
                    metadata={"start_row": 2, "end_row": 2 + len(chunk_rows) - 1}
                ))
                chunk_rows = [row]
                current_size = len(header_text) + row_size + len("[CSV]\n")
            else:
                chunk_rows.append(row)
                current_size += row_size

        if chunk_rows:
            start = 2 + len(data_rows) - len(chunk_rows)
            chunk_text = TextSplitter._format_csv_chunk(
                header_text, chunk_rows,
                start_row=start, end_row=start + len(chunk_rows) - 1
            )
            chunks.append(Chunk(
                text=chunk_text,
                chunk_id=f"csv_rows_{start}-{start + len(chunk_rows) - 1}",
                chunk_type="csv_rows",
                metadata={"start_row": start, "end_row": start + len(chunk_rows) - 1}
            ))

        return chunks if chunks else [Chunk(text="[CSVからテキストを抽出できませんでした]", chunk_id="error", chunk_type="error")]

    @staticmethod
    def _format_csv_chunk(header_text: str, rows: List[List[str]], start_row: int, end_row: int) -> str:
        """CSVチャンクをフォーマット"""
        lines = [f"[CSV / 行 {start_row}-{end_row}]"]
        lines.append(header_text)
        lines.append("-" * min(len(header_text), 60))
        for row in rows:
            lines.append(",".join(row))
        return '\n'.join(lines)

    # ============================================================
    # 共通ユーティリティ
    # ============================================================
    @staticmethod
    def _is_heading(line: str) -> bool:
        """行が見出しかどうかを判定"""
        for pattern in TextSplitter.HEADING_PATTERNS:
            if re.match(pattern, line):
                return True
        # 大文字のみの短い行も見出しの可能性
        if len(line) <= 50 and line.isupper():
            return True
        return False

    @staticmethod
    def _split_text_semantic(text: str, chunk_id_prefix: str = "chunk", chunk_type: str = "topic") -> List[Chunk]:
        """テキストを意味的にチャンク分割（見出し→段落→文の優先度で分割）"""
        if len(text) <= TextSplitter.MAX_CHUNK_SIZE:
            return [Chunk(text=text, chunk_id=f"{chunk_id_prefix}_1", chunk_type=chunk_type)]

        # トピック単位で分割
        topics = TextSplitter._split_by_structure(text)

        chunks = []
        current_text = ""
        chunk_num = 1

        for topic in topics:
            # トピックが大きすぎる場合は文単位で分割
            if len(topic) > TextSplitter.MAX_CHUNK_SIZE:
                sentences = TextSplitter._split_into_sentences(topic)
                for sentence in sentences:
                    test_text = current_text + ("\n" if current_text else "") + sentence
                    if len(test_text) > TextSplitter.MAX_CHUNK_SIZE and current_text:
                        chunks.append(Chunk(
                            text=current_text,
                            chunk_id=f"{chunk_id_prefix}_{chunk_num}",
                            chunk_type=chunk_type
                        ))
                        chunk_num += 1
                        # オーバーラップ
                        overlap = current_text[-TextSplitter.OVERLAP_SIZE:] if len(current_text) > TextSplitter.OVERLAP_SIZE else current_text
                        current_text = "... " + overlap + "\n" + sentence
                    else:
                        current_text = test_text
            else:
                test_text = current_text + ("\n\n" if current_text else "") + topic
                if len(test_text) > TextSplitter.MAX_CHUNK_SIZE and current_text:
                    chunks.append(Chunk(
                        text=current_text,
                        chunk_id=f"{chunk_id_prefix}_{chunk_num}",
                        chunk_type=chunk_type
                    ))
                    chunk_num += 1
                    overlap = current_text[-TextSplitter.OVERLAP_SIZE:] if len(current_text) > TextSplitter.OVERLAP_SIZE else current_text
                    current_text = "... " + overlap + "\n\n" + topic
                else:
                    current_text = test_text

        # 残り
        if current_text:
            if len(current_text) < TextSplitter.MIN_CHUNK_SIZE and chunks:
                chunks[-1] = Chunk(
                    text=chunks[-1].text + "\n\n" + current_text,
                    chunk_id=chunks[-1].chunk_id,
                    chunk_type=chunks[-1].chunk_type,
                    metadata=chunks[-1].metadata
                )
            else:
                chunks.append(Chunk(
                    text=current_text,
                    chunk_id=f"{chunk_id_prefix}_{chunk_num}",
                    chunk_type=chunk_type
                ))

        return chunks

    @staticmethod
    def _split_by_structure(text: str) -> List[str]:
        """テキストを見出し・空行でセクション分割"""
        # 見出しパターンで分割
        heading_pattern = r'\n(?=(?:' + '|'.join(TextSplitter.HEADING_PATTERNS) + '))'
        sections = re.split(heading_pattern, text)

        result = []
        for section in sections:
            section = section.strip()
            if not section:
                continue

            # セクションが大きすぎる場合は空行で分割
            if len(section) > TextSplitter.MAX_CHUNK_SIZE:
                paragraphs = re.split(r'\n\s*\n', section)
                for para in paragraphs:
                    para = para.strip()
                    if para:
                        result.append(para)
            else:
                result.append(section)

        return result

    @staticmethod
    def _split_into_sentences(text: str) -> List[str]:
        """テキストを文単位で分割（日本語・英語対応）"""
        sentences = re.split(r'(?<=[。！？.!?])\s*', text)
        return [s.strip() for s in sentences if s.strip()]

    @staticmethod
    def _build_chunks_from_sections(sections: List[Tuple[int, str]], chunk_type: str) -> List[Chunk]:
        """(page_num, text) のリストからチャンクを構築"""
        chunks = []
        current_text = ""
        current_page = 1
        chunk_num = 1

        for page_num, text in sections:
            header = f"[ページ {page_num}]\n\n"
            section_text = header + text

            test_text = current_text + ("\n\n" if current_text else "") + section_text

            if len(test_text) > TextSplitter.MAX_CHUNK_SIZE and current_text:
                chunks.append(Chunk(
                    text=current_text,
                    chunk_id=f"chunk_{chunk_num}",
                    chunk_type=chunk_type,
                    metadata={"page": current_page}
                ))
                chunk_num += 1
                # オーバーラップ
                overlap = current_text[-TextSplitter.OVERLAP_SIZE:] if len(current_text) > TextSplitter.OVERLAP_SIZE else ""
                current_text = ("... " + overlap + "\n\n" if overlap else "") + section_text
                current_page = page_num
            else:
                current_text = test_text
                if not current_text.startswith("[ページ"):
                    current_page = page_num

        if current_text:
            if len(current_text) < TextSplitter.MIN_CHUNK_SIZE and chunks:
                chunks[-1] = Chunk(
                    text=chunks[-1].text + "\n\n" + current_text,
                    chunk_id=chunks[-1].chunk_id,
                    chunk_type=chunks[-1].chunk_type,
                    metadata=chunks[-1].metadata
                )
            else:
                chunks.append(Chunk(
                    text=current_text,
                    chunk_id=f"chunk_{chunk_num}",
                    chunk_type=chunk_type,
                    metadata={"page": current_page}
                ))

        return chunks

    @staticmethod
    def _merge_small_chunks(chunks: List[Chunk]) -> List[Chunk]:
        """小さすぎるチャンクを前のチャンクに結合"""
        if not chunks:
            return chunks

        merged = [chunks[0]]
        for chunk in chunks[1:]:
            if len(merged[-1].text) + len(chunk.text) < TextSplitter.MAX_CHUNK_SIZE and len(chunk.text) < TextSplitter.MIN_CHUNK_SIZE:
                merged[-1] = Chunk(
                    text=merged[-1].text + "\n\n" + chunk.text,
                    chunk_id=merged[-1].chunk_id,
                    chunk_type=merged[-1].chunk_type,
                    metadata=merged[-1].metadata
                )
            else:
                merged.append(chunk)
        return merged
